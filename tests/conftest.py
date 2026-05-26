from __future__ import annotations

import io
import os
import uuid
from dataclasses import dataclass

import pytest
import requests
from docx import Document


HEALTH_PATH = "/health"


@dataclass(frozen=True)
class AuthContext:
    username: str
    password: str
    email: str
    access_token: str


def _resolve_base_url() -> str:
    env_url = os.getenv("TEST_API_BASE_URL")
    candidates = [env_url] if env_url else [
        "http://127.0.0.1:8000",
        "http://127.0.0.1:8001",
    ]

    for base_url in candidates:
        if not base_url:
            continue
        try:
            response = requests.get(f"{base_url}{HEALTH_PATH}", timeout=3)
        except requests.RequestException:
            continue
        if response.status_code == 200:
            return base_url

    raise RuntimeError(
        "Unable to reach the API. Start the service first or set TEST_API_BASE_URL."
    )


def _register_and_login(client: requests.Session, base_url: str) -> AuthContext:
    suffix = uuid.uuid4().hex[:10]
    username = f"pytest_{suffix}"
    password = "secret123"
    email = f"{username}@example.com"

    register_response = client.post(
        f"{base_url}/v1/auth/register",
        json={
            "username": username,
            "email": email,
            "full_name": "Pytest Runner",
            "password": password,
        },
        timeout=15,
    )
    register_response.raise_for_status()

    login_response = client.post(
        f"{base_url}/v1/auth/login",
        json={"username": username, "password": password},
        timeout=15,
    )
    login_response.raise_for_status()

    return AuthContext(
        username=username,
        password=password,
        email=email,
        access_token=login_response.json()["access_token"],
    )


def build_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("DOCX 测试文档", level=1)
    document.add_paragraph("这是一个用于验证 DOCX 转 Markdown 的测试段落。")
    document.add_paragraph("第一项", style="List Bullet")
    document.add_paragraph("第二项", style="List Bullet")
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture(scope="session")
def api_base_url() -> str:
    return _resolve_base_url()


@pytest.fixture
def client() -> requests.Session:
    session = requests.Session()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def auth_context(client: requests.Session, api_base_url: str) -> AuthContext:
    return _register_and_login(client, api_base_url)


@pytest.fixture
def auth_headers(auth_context: AuthContext) -> dict[str, str]:
    return {"Authorization": f"Bearer {auth_context.access_token}"}


@pytest.fixture
def docx_bytes() -> bytes:
    return build_docx_bytes()


@pytest.fixture
def uploaded_markdown_doc(
    client: requests.Session,
    api_base_url: str,
    auth_headers: dict[str, str],
) -> dict[str, str]:
    response = client.post(
        f"{api_base_url}/v1/docs/upload",
        data={
            "title": "Pytest Markdown",
            "content": (
                "# 项目需求文档\n\n"
                "## 技术架构\n\n"
                "系统采用 FastAPI + PostgreSQL 架构。\n\n"
                "## 核心功能\n\n"
                "支持对话式编辑。"
            ),
        },
        headers=auth_headers,
        timeout=20,
    )
    response.raise_for_status()
    return response.json()


@pytest.fixture
def uploaded_bulk_doc(
    client: requests.Session,
    api_base_url: str,
    auth_headers: dict[str, str],
) -> dict[str, str]:
    response = client.post(
        f"{api_base_url}/v1/docs/upload",
        data={
            "title": "Pytest Bulk",
            "content": (
                "# 批量替换测试\n\n"
                "旧词 要被替换。\n\n"
                "## 范围\n\n"
                "这里还有旧词。\n\n"
                "最后再来一个旧词。"
            ),
        },
        headers=auth_headers,
        timeout=20,
    )
    response.raise_for_status()
    return response.json()
